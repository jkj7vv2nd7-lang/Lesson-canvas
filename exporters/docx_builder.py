"""
Markdown形式のAI生成テキストを、きれいに表組みされたWord文書に変換する。

旧 app.py の create_word_file() をベースに、
「単元計画」以外の成果物タイプ（本時案・ワークシート・小テスト等）でも
使えるよう、タイトル・メタ情報を汎用化した。
"""

from __future__ import annotations

import io
import re

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor


def clean_html_tags(text: str) -> str:
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    return re.sub(r"</?[^>]+>", "", text)


def sanitize_filename(name: str) -> str:
    return re.sub(r'[\\/*?:"<>|]', "", name)


def _set_cell_background(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def _is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def _is_table_separator(line: str) -> bool:
    stripped = line.strip()
    return re.match(r"^\|[\s:\-|_]+\|$", stripped) is not None


def _flush_code_block(doc, code_buffer: list[str]) -> None:
    if not code_buffer:
        return
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.text = "\n".join(code_buffer)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(10)
    _set_cell_background(cell, "F2F2F2")
    doc.add_paragraph()


def markdown_to_docx(
    *,
    title: str,
    subtitle: str,
    content_text: str,
    footer_note: str = "",
) -> io.BytesIO:
    """
    title: 文書タイトル（例: "がい数 単元指導計画" / "第1時 本時案"）
    subtitle: タイトル下に表示するメタ情報（例: "学年: 4年 | 教科: 算数 | 総時数: 6時間"）
    content_text: AIが生成したMarkdownテキスト
    footer_note: 文末に入れる注記（著作権配慮の文言など）
    """
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    title_p = doc.add_paragraph(title)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title_p.runs:
        run = title_p.runs[0]
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(43, 88, 118)

    if subtitle:
        meta_p = doc.add_paragraph(subtitle)
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("-" * 50)

    cleaned = clean_html_tags(content_text)
    lines = cleaned.split("\n")

    in_code_block = False
    code_buffer: list[str] = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                _flush_code_block(doc, code_buffer)
                in_code_block = False
                code_buffer = []
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        if _is_table_line(line):
            table_lines = []
            while i < len(lines) and _is_table_line(lines[i]):
                table_lines.append(lines[i].strip())
                i += 1

            valid_rows = [r for r in table_lines if not _is_table_separator(r)]

            if valid_rows:
                rows_data = []
                for r in valid_rows:
                    content = r[1:-1]
                    cells = [cell.strip() for cell in content.split("|")]
                    rows_data.append(cells)

                col_count = max(len(r) for r in rows_data)
                for row in rows_data:
                    while len(row) < col_count:
                        row.append("")

                table = doc.add_table(rows=len(rows_data), cols=col_count)
                table.style = "Table Grid"
                table.autofit = False

                section = doc.sections[0]
                total_width = section.page_width - section.left_margin - section.right_margin
                if col_count > 0:
                    col_width = int(total_width / col_count)
                    for col in table.columns:
                        col.width = col_width

                for row_idx, r_data in enumerate(rows_data):
                    for col_idx, cell_value in enumerate(r_data):
                        if col_idx < col_count:
                            cell = table.cell(row_idx, col_idx)
                            cell.text = cell_value
                            for p in cell.paragraphs:
                                p.paragraph_format.line_spacing = Pt(12)
                                for run in p.runs:
                                    run.font.size = Pt(10)
                            if row_idx == 0:
                                _set_cell_background(cell, "2B5876")
                                for p in cell.paragraphs:
                                    for run in p.runs:
                                        run.font.bold = True
                                        run.font.color.rgb = RGBColor(255, 255, 255)
            doc.add_paragraph()
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped.replace("# ", ""), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped.replace("## ", ""), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped.replace("### ", ""), level=3)
        elif stripped.startswith("- ") or stripped.startswith("・"):
            doc.add_paragraph(stripped, style="List Bullet")
        elif stripped:
            doc.add_paragraph(stripped)

        i += 1

    if in_code_block and code_buffer:
        _flush_code_block(doc, code_buffer)

    if footer_note:
        doc.add_paragraph()
        note_p = doc.add_paragraph(footer_note)
        for run in note_p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(120, 120, 120)
            run.font.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
